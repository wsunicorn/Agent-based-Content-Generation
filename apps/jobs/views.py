"""API views for the jobs app."""
import io
import re
import textwrap

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.utils import timezone
from rest_framework import status
from rest_framework.decorators import api_view
from rest_framework.response import Response

from .models import Artifact, Job
from .serializers import (
    ArtifactSerializer,
    JobCreateSerializer,
    JobDetailSerializer,
    JobListSerializer,
)
from .tasks import run_pipeline


def _normalise_create_payload(request):
    if hasattr(request.data, "lists"):
        data = {
            key: values[-1] if len(values) == 1 else values
            for key, values in request.data.lists()
        }
    else:
        data = dict(request.data)

    keywords = data.get("keywords")
    if isinstance(keywords, str):
        data["keywords"] = [kw.strip() for kw in keywords.split(",") if kw.strip()]
    elif isinstance(keywords, list) and len(keywords) == 1 and isinstance(keywords[0], str):
        data["keywords"] = [kw.strip() for kw in keywords[0].split(",") if kw.strip()]

    return data


@api_view(["GET", "POST"])
def job_list_create(request):
    if request.method == "GET":
        jobs = Job.objects.all()
        serializer = JobListSerializer(jobs, many=True)
        return Response(serializer.data)

    payload = _normalise_create_payload(request)
    serializer = JobCreateSerializer(data=payload)
    serializer.is_valid(raise_exception=True)
    job = serializer.save()

    import uuid
    from django.db import transaction

    # Predetermine task ID to prevent Django-Celery transaction race condition
    task_id = str(uuid.uuid4())
    job.celery_task_id = task_id
    job.status = Job.Status.RUNNING
    job.save(update_fields=["celery_task_id", "status"])

    # Dispatch Celery task only after transaction has successfully committed
    transaction.on_commit(lambda: run_pipeline.apply_async(args=[str(job.id)], task_id=task_id))

    return Response(JobDetailSerializer(job).data, status=status.HTTP_201_CREATED)


@api_view(["GET", "DELETE"])
def job_detail(request, pk):
    job = get_object_or_404(Job, pk=pk)
    if request.method == "DELETE":
        # Revoke Celery task if still running
        if job.celery_task_id and job.status in (Job.Status.PENDING, Job.Status.RUNNING):
            try:
                from celery.app.control import Control
                from config.celery import app as celery_app
                celery_app.control.revoke(job.celery_task_id, terminate=True)
            except Exception:
                pass
        job.delete()
        return Response(status=status.HTTP_204_NO_CONTENT)
    serializer = JobDetailSerializer(job)
    return Response(serializer.data)


@api_view(["PATCH"])
def job_update_content(request, pk):
    """Update the final_content artifact text (manual edit before export)."""
    job = get_object_or_404(Job, pk=pk)
    new_text = request.data.get("content_text", "").strip()
    if not new_text:
        return Response({"detail": "content_text is required."}, status=status.HTTP_400_BAD_REQUEST)

    artifact = (
        job.artifacts.filter(artifact_type=Artifact.ArtifactType.FINAL_CONTENT)
        .order_by("-version").first()
    )
    if not artifact:
        return Response({"detail": "Final content artifact not found."}, status=status.HTTP_404_NOT_FOUND)

    artifact.content_text = new_text
    artifact.word_count = len(new_text.split())
    artifact.save(update_fields=["content_text", "word_count"])
    return Response({"detail": "Content updated.", "word_count": artifact.word_count})


@api_view(["GET"])
def job_artifact(request, pk, artifact_type):
    job = get_object_or_404(Job, pk=pk)
    artifact = (
        job.artifacts.filter(artifact_type=artifact_type).order_by("-version").first()
    )
    if not artifact:
        return Response(
            {"detail": "Artifact not found."}, status=status.HTTP_404_NOT_FOUND
        )
    return Response(ArtifactSerializer(artifact).data)


@api_view(["GET"])
def job_evidence(request, pk):
    job = get_object_or_404(Job, pk=pk)
    source_artifact = (
        job.artifacts.filter(artifact_type=Artifact.ArtifactType.SOURCE_DOCUMENTS)
        .order_by("-version", "-created_at")
        .first()
    )
    image_artifact = (
        job.artifacts.filter(artifact_type=Artifact.ArtifactType.IMAGE_ASSETS)
        .order_by("-version", "-created_at")
        .first()
    )
    outline_artifact = (
        job.artifacts.filter(artifact_type=Artifact.ArtifactType.OUTLINE)
        .order_by("-version", "-created_at")
        .first()
    )
    sources = []
    if source_artifact and isinstance(source_artifact.content_json, dict):
        sources = source_artifact.content_json.get("sources") or []
    images = []
    if image_artifact and isinstance(image_artifact.content_json, dict):
        images = image_artifact.content_json.get("image_assets") or []
    outline = []
    if outline_artifact and isinstance(outline_artifact.content_json, dict):
        outline = outline_artifact.content_json.get("sections") or []
    return Response(
        {
            "sources": sources if isinstance(sources, list) else [],
            "images": images if isinstance(images, list) else [],
            "outline": outline if isinstance(outline, list) else [],
        }
    )


@api_view(["POST"])
def job_approve_outline(request, pk):
    job = get_object_or_404(Job, pk=pk)
    if not job.pipeline_state:
        return Response(
            {"detail": "No paused pipeline checkpoint is available for this job."},
            status=status.HTTP_400_BAD_REQUEST,
        )
    if job.status != Job.Status.PAUSED:
        return Response(
            {"detail": "Job is not waiting for outline approval."},
            status=status.HTTP_400_BAD_REQUEST,
        )
    sections = _normalise_outline_sections(request.data.get("sections") or [])
    if not sections:
        latest = (
            job.artifacts.filter(artifact_type=Artifact.ArtifactType.OUTLINE)
            .order_by("-version", "-created_at")
            .first()
        )
        if latest and isinstance(latest.content_json, dict):
            sections = _normalise_outline_sections(latest.content_json.get("sections") or [])
    if not sections:
        return Response({"detail": "Outline sections are required."}, status=status.HTTP_400_BAD_REQUEST)

    job.approved_outline = sections
    job.outline_approved_at = timezone.now()
    checkpoint = dict(job.pipeline_state or {})
    checkpoint["sections"] = sections
    checkpoint["outline_approved"] = True
    checkpoint["completed"] = False
    job.pipeline_state = checkpoint
    job.status = Job.Status.RUNNING

    job.save(update_fields=[
        "approved_outline",
        "outline_approved_at",
        "pipeline_state",
        "status",
    ])
    Artifact.objects.create(
        job=job,
        artifact_type=Artifact.ArtifactType.OUTLINE,
        content_json={"sections": sections, "approved": True},
        version=_next_artifact_version(job, Artifact.ArtifactType.OUTLINE),
    )
    import uuid
    from django.db import transaction

    # Predetermine task ID to prevent Django-Celery transaction race condition
    task_id = str(uuid.uuid4())
    job.celery_task_id = task_id
    job.save(update_fields=["celery_task_id"])

    # Dispatch Celery task only after transaction has successfully committed
    transaction.on_commit(lambda: run_pipeline.apply_async(args=[str(job.id)], task_id=task_id))

    return Response({"detail": "Outline approved.", "task_id": task_id, "sections": sections})


@api_view(["POST"])
def job_regenerate_section(request, pk, section_id):
    job = get_object_or_404(Job, pk=pk)
    if job.status != Job.Status.COMPLETED:
        return Response(
            {"detail": "Section regeneration is available after a completed job."},
            status=status.HTTP_400_BAD_REQUEST,
        )
    if not job.pipeline_state:
        return Response(
            {"detail": "No pipeline checkpoint available for this job."},
            status=status.HTTP_400_BAD_REQUEST,
        )
    checkpoint = dict(job.pipeline_state)
    sections = checkpoint.get("sections") or job.approved_outline or []
    if not sections:
        return Response({"detail": "No outline found for regeneration."}, status=status.HTTP_400_BAD_REQUEST)

    checkpoint["sections"] = _normalise_outline_sections(sections)
    checkpoint["outline_approved"] = True
    checkpoint["revision_target_section_ids"] = [int(section_id)]
    checkpoint["revision_instructions"] = (
        request.data.get("instructions")
        or f"Regenerate section {section_id} while preserving the rest of the article."
    )
    checkpoint["revision_count"] = int(checkpoint.get("revision_count") or 0) + 1
    checkpoint["completed"] = False
    checkpoint["final_content"] = ""

    job.pipeline_state = checkpoint
    job.approved_outline = checkpoint["sections"]
    job.outline_approved_at = job.outline_approved_at or timezone.now()
    job.status = Job.Status.RUNNING
    job.save(update_fields=[
        "pipeline_state",
        "approved_outline",
        "outline_approved_at",
        "status",
    ])
    import uuid
    from django.db import transaction

    # Predetermine task ID to prevent Django-Celery transaction race condition
    task_id = str(uuid.uuid4())
    job.celery_task_id = task_id
    job.save(update_fields=["celery_task_id"])

    # Dispatch Celery task only after transaction has successfully committed
    transaction.on_commit(lambda: run_pipeline.apply_async(args=[str(job.id)], task_id=task_id))

    return Response({"detail": "Section regeneration started.", "task_id": task_id})


def _normalise_outline_sections(raw_sections):
    sections = []
    for item in raw_sections:
        if not isinstance(item, dict):
            continue
        heading = str(item.get("heading") or "").strip()
        brief = str(item.get("brief") or "").strip()
        if not heading:
            continue
        key_points = item.get("key_points") or []
        if isinstance(key_points, str):
            key_points = [line.strip("- ").strip() for line in key_points.splitlines()]
        sections.append(
            {
                "heading": heading[:180],
                "level": int(item.get("level") or 1),
                "brief": brief[:1000],
                "key_points": [str(point).strip() for point in key_points if str(point).strip()][:8],
                "template_role": str(item.get("template_role") or "").strip()[:120],
            }
        )
    return sections


def _next_artifact_version(job, artifact_type):
    latest = (
        Artifact.objects.filter(job=job, artifact_type=artifact_type)
        .order_by("-version")
        .values_list("version", flat=True)
        .first()
    )
    return int(latest or 0) + 1


@api_view(["POST"])
def job_cancel(request, pk):
    job = get_object_or_404(Job, pk=pk)
    if job.status not in (Job.Status.PENDING, Job.Status.RUNNING, Job.Status.PAUSED):
        return Response(
            {"detail": "Job cannot be cancelled in its current state."},
            status=status.HTTP_400_BAD_REQUEST,
        )
    job.status = Job.Status.CANCELLED
    job.save(update_fields=["status"])
    return Response({"detail": "Job cancelled."})


# ---------------------------------------------------------------------------
# Export endpoint
# ---------------------------------------------------------------------------

@api_view(["GET"])
def job_export(request, pk):
    """
    Export article as Markdown, HTML, or DOCX.
    GET /api/jobs/<id>/export/?format=markdown|html|docx
    """
    job = get_object_or_404(Job, pk=pk)
    if job.status != Job.Status.COMPLETED:
        return Response(
            {"detail": "Job is not completed yet."},
            status=status.HTTP_400_BAD_REQUEST,
        )

    fmt = (
        request.query_params.get("type")
        or request.query_params.get("format")
        or "markdown"
    ).lower()
    if fmt not in ("markdown", "html", "docx"):
        return Response(
            {"detail": "format must be one of: markdown, html, docx"},
            status=status.HTTP_400_BAD_REQUEST,
        )

    # Fetch artifacts
    final_artifact = (
        job.artifacts.filter(artifact_type=Artifact.ArtifactType.FINAL_CONTENT)
        .order_by("-version").first()
    )
    if not final_artifact or not final_artifact.content_text:
        return Response(
            {"detail": "Final content artifact not found."},
            status=status.HTTP_404_NOT_FOUND,
        )

    seo_artifact = (
        job.artifacts.filter(artifact_type=Artifact.ArtifactType.SEO_METADATA)
        .order_by("-version").first()
    )
    qa_artifact = (
        job.artifacts.filter(artifact_type=Artifact.ArtifactType.QA_REPORT)
        .order_by("-version").first()
    )
    image_artifact = (
        job.artifacts.filter(artifact_type=Artifact.ArtifactType.IMAGE_ASSETS)
        .order_by("-version").first()
    )

    content = final_artifact.content_text
    seo = seo_artifact.content_json if seo_artifact else {}
    qa_score = (qa_artifact.content_json or {}).get("overall_score", "N/A") if qa_artifact else "N/A"
    image_refs = _image_references(image_artifact)

    slug = seo.get("slug") or re.sub(r"[^a-z0-9]+", "-", job.title.lower())[:60]
    filename_base = slug.strip("-")

    if fmt == "markdown":
        md = _build_markdown(job, content, seo, qa_score, image_refs)
        response = HttpResponse(md, content_type="text/markdown; charset=utf-8")
        response["Content-Disposition"] = f'attachment; filename="{filename_base}.md"'
        return response

    elif fmt == "html":
        html = _build_html(job, content, seo, qa_score, image_refs)
        response = HttpResponse(html, content_type="text/html; charset=utf-8")
        response["Content-Disposition"] = f'attachment; filename="{filename_base}.html"'
        return response

    else:  # docx
        doc_bytes = _build_docx(job, content, seo, qa_score, image_refs)
        response = HttpResponse(
            doc_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename_base}.docx"'
        return response


def _image_references(artifact):
    if not artifact or not isinstance(artifact.content_json, dict):
        return []
    refs = artifact.content_json.get("image_assets") or []
    return refs if isinstance(refs, list) else []


def _image_source_markdown(image_refs):
    if not image_refs:
        return ""
    lines = ["", "", "## Image Sources"]
    for item in image_refs:
        if not isinstance(item, dict):
            continue
        title = item.get("title") or "Image"
        url = item.get("url") or ""
        source = item.get("source_url") or url
        alt_text = item.get("alt_text") or title
        
        if url:
            lines.append(f"![{alt_text}]({url})")
        
        caption_parts = []
        if item.get("caption"):
            caption_parts.append(f"**{item.get('caption')}**")
        else:
            caption_parts.append(f"**{title}**")
            
        if item.get("license"):
            caption_parts.append(f"License: {item.get('license')}")
        if item.get("attribution"):
            caption_parts.append(f"Attribution: {item.get('attribution')}")
        if source:
            caption_parts.append(f"[Source link]({source})")
            
        lines.append(" - ".join(caption_parts))
        lines.append("")
    return "\n".join(lines) + "\n"


def _build_markdown(job, content, seo, qa_score, image_refs=None):
    meta_title = seo.get("meta_title", job.title)
    meta_desc  = seo.get("meta_description", "")
    focus_kw   = seo.get("focus_keyword", "")
    date_str   = (job.completed_at or timezone.now()).strftime("%Y-%m-%d")

    header = textwrap.dedent(f"""\
        ---
        title: "{meta_title}"
        description: "{meta_desc}"
        focus_keyword: "{focus_kw}"
        content_type: "{job.content_type}"
        domain: "{getattr(job, 'domain', '')}"
        audience: "{getattr(job, 'audience', '')}"
        tone: "{getattr(job, 'tone', '')}"
        word_count: {job.artifacts.filter(artifact_type=Artifact.ArtifactType.FINAL_CONTENT).values_list('word_count', flat=True).first() or 'N/A'}
        qa_score: {qa_score}
        generated: "{date_str}"
        ---

        """)
    return header + content + _image_source_markdown(image_refs or [])


def _build_html(job, content, seo, qa_score, image_refs=None):
    import markdown2
    meta_title = seo.get("meta_title", job.title)
    meta_desc  = seo.get("meta_description", "")
    focus_kw   = seo.get("focus_keyword", "")
    date_str   = (job.completed_at or timezone.now()).strftime("%B %d, %Y")

    body_html = markdown2.markdown(
        content + _image_source_markdown(image_refs or []),
        extras=["fenced-code-blocks", "tables", "header-ids"],
    )

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <meta name="description" content="{_esc(meta_desc)}">
  <meta name="keywords" content="{_esc(focus_kw)}">
  <title>{_esc(meta_title)}</title>
  <style>
    :root {{ --text:#1a1a2e; --muted:#6b7280; --accent:#7c6cf8; --border:#e5e7eb; }}
    * {{ box-sizing:border-box; margin:0; padding:0; }}
    body {{ font-family:'Georgia',serif; color:var(--text); background:#fff; line-height:1.8; }}
    .wrapper {{ max-width:740px; margin:0 auto; padding:60px 24px; }}
    .meta {{ font-family:'Inter',sans-serif; font-size:.8rem; color:var(--muted); margin-bottom:40px; }}
    .meta span {{ margin-right:20px; }}
    .meta .badge {{ background:var(--accent); color:#fff; padding:2px 10px; border-radius:99px; font-weight:600; }}
    h1 {{ font-size:2.2rem; line-height:1.25; margin-bottom:12px; letter-spacing:-.02em; }}
    h2 {{ font-size:1.4rem; margin:2rem 0 .75rem; padding-top:.5rem; border-top:1px solid var(--border); }}
    h3 {{ font-size:1.15rem; margin:1.6rem 0 .5rem; }}
    p  {{ margin-bottom:1.2rem; }}
    ul,ol {{ margin:0 0 1.2rem 1.4rem; }}
    li {{ margin-bottom:.4rem; }}
    code {{ background:#f3f4f6; padding:2px 6px; border-radius:4px; font-size:.88em; font-family:monospace; }}
    pre  {{ background:#f3f4f6; padding:16px; border-radius:8px; overflow-x:auto; margin-bottom:1.2rem; }}
    blockquote {{ border-left:3px solid var(--accent); padding-left:16px; color:var(--muted); margin:1.2rem 0; }}
    .footer {{ margin-top:60px; padding-top:20px; border-top:1px solid var(--border); font-family:'Inter',sans-serif; font-size:.75rem; color:var(--muted); }}
  </style>
</head>
<body>
<div class="wrapper">
  <div class="meta">
    <span>{_esc(job.content_type.replace('_',' ').title())}</span>
    <span>{date_str}</span>
    <span class="badge">QA {qa_score}/100</span>
  </div>
  {body_html}
  <div class="footer">
    Generated by Domain LLM Assistant · {_esc(job.title)}
  </div>
</div>
</body>
</html>"""


def _build_docx(job, content, seo, qa_score, image_refs=None):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page margins
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

    meta_title = seo.get("meta_title", job.title)
    focus_kw   = seo.get("focus_keyword", "")
    date_str   = (job.completed_at or timezone.now()).strftime("%B %d, %Y")

    # Document title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(meta_title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    # Subtitle / meta line
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_para.add_run(
        f"{job.content_type.replace('_',' ').title()}  ·  {date_str}  ·  QA Score: {qa_score}/100"
    )
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_paragraph()  # spacer

    # Parse & write content
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:], style="Heading 3")
        elif stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:], style="Heading 2")
        elif stripped.startswith("# "):
            p = doc.add_paragraph(stripped[2:], style="Heading 1")
        elif stripped.startswith(("- ", "* ")):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", stripped), style="List Number")
        else:
            # Handle inline bold (**text**)
            para = doc.add_paragraph()
            _add_inline_formatted(para, stripped)

    if image_refs:
        doc.add_paragraph()
        doc.add_paragraph("Image sources", style="Heading 2")
        for item in image_refs:
            label = item.get("caption") or item.get("alt_text") or item.get("title") or "Image"
            source = item.get("source_url") or item.get("url") or ""
            suffix = f" Source: {source}" if source else ""
            doc.add_paragraph(f"{label}{suffix}", style="List Bullet")

    # Footer metadata
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(
        f"Focus keyword: {focus_kw}  ·  Generated by Domain LLM Assistant"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline_formatted(para, text):
    """Write a paragraph handling **bold** inline markers."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        else:
            para.add_run(part)


def _esc(s):
    """Minimal HTML escaping for attributes."""
    return str(s).replace("&", "&amp;").replace('"', "&quot;").replace("<", "&lt;").replace(">", "&gt;")


# ---------------------------------------------------------------------------
# Analytics endpoint
# ---------------------------------------------------------------------------

@api_view(["GET"])
def analytics_summary(request):
    """
    GET /api/analytics/
    Returns aggregate stats about all jobs.
    """
    from django.db.models import Avg, Count, Sum, F, ExpressionWrapper, DurationField
    from django.db.models.functions import TruncDate

    jobs = Job.objects.all()
    total       = jobs.count()
    completed   = jobs.filter(status=Job.Status.COMPLETED).count()
    failed      = jobs.filter(status=Job.Status.FAILED).count()
    running     = jobs.filter(status=Job.Status.RUNNING).count()

    # Average QA score from QA_REPORT artifacts
    qa_scores = list(
        Artifact.objects.filter(artifact_type=Artifact.ArtifactType.QA_REPORT)
        .exclude(content_json=None)
        .values_list("content_json", flat=True)
    )
    valid_scores = [
        float(q["overall_score"]) for q in qa_scores
        if isinstance(q, dict) and "overall_score" in q
    ]
    avg_qa = round(sum(valid_scores) / len(valid_scores), 1) if valid_scores else None

    # Average duration (seconds) for completed jobs
    completed_jobs = jobs.filter(
        status=Job.Status.COMPLETED,
        started_at__isnull=False,
        completed_at__isnull=False,
    )
    durations = [
        (j.completed_at - j.started_at).total_seconds()
        for j in completed_jobs
        if j.completed_at and j.started_at
    ]
    avg_duration = round(sum(durations) / len(durations), 1) if durations else None

    # Total LLM calls
    total_llm_calls = jobs.aggregate(total=Sum("llm_calls_count"))["total"] or 0
    llm_calls_by_provider = {}
    for usage in jobs.values_list("llm_usage_by_provider", flat=True):
        if not isinstance(usage, dict):
            continue
        for provider, values in usage.items():
            if isinstance(values, dict):
                llm_calls_by_provider[provider] = (
                    llm_calls_by_provider.get(provider, 0)
                    + int(values.get("calls") or 0)
                )

    # Jobs per day (last 7 days)
    from django.utils import timezone as tz
    from datetime import timedelta
    week_ago = tz.now() - timedelta(days=7)
    daily = (
        jobs.filter(created_at__gte=week_ago)
        .values(date=TruncDate("created_at"))
        .annotate(count=Count("id"))
        .order_by("date")
    )

    # Per-job stats for chart
    recent_jobs = (
        jobs.filter(status=Job.Status.COMPLETED).order_by("-created_at")[:20]
    )
    job_stats = []
    for j in recent_jobs:
        qa_art = j.artifacts.filter(artifact_type=Artifact.ArtifactType.QA_REPORT).first()
        score = None
        if qa_art and isinstance(qa_art.content_json, dict):
            score = qa_art.content_json.get("overall_score")
        job_stats.append({
            "id": str(j.id),
            "title": j.title[:50],
            "content_type": j.content_type,
            "llm_calls": j.llm_calls_count,
            "duration": j.duration_seconds,
            "qa_score": score,
            "created_at": j.created_at.isoformat(),
        })

    return Response({
        "totals": {
            "total": total,
            "completed": completed,
            "failed": failed,
            "running": running,
            "success_rate": round(completed / total * 100, 1) if total else 0,
        },
        "averages": {
            "qa_score": avg_qa,
            "duration_seconds": avg_duration,
            "llm_calls": round(total_llm_calls / total, 1) if total else 0,
        },
        "totals_llm": total_llm_calls,
        "llm_calls_by_provider": llm_calls_by_provider,
        "daily_jobs": [
            {"date": str(d["date"]), "count": d["count"]} for d in daily
        ],
        "recent_jobs": job_stats,
    })


@api_view(["GET"])
def health_check(request):
    """GET /api/health/ — liveness probe."""
    from django.db import connection
    from django.core.cache import cache
    import datetime

    db_ok = False
    try:
        connection.ensure_connection()
        db_ok = True
    except Exception:
        pass

    redis_ok = False
    try:
        cache.set("_health", "1", timeout=5)
        redis_ok = cache.get("_health") == "1"
    except Exception:
        pass

    worker_ok = False
    worker_count = 0
    try:
        # Try control.ping first (works with some pool types)
        from config.celery import app as celery_app
        replies = celery_app.control.ping(timeout=1.0) or []
        worker_count = len(replies)
        worker_ok = worker_count > 0
    except Exception:
        pass

    # Fallback: check Redis for active worker registrations
    if not worker_ok:
        try:
            import redis as redis_lib
            from django.conf import settings as django_settings

            r = redis_lib.from_url(django_settings.CELERY_BROKER_URL)
            # Check if there are workers registered via kombu bindings
            bindings = r.smembers("_kombu.binding.celery")
            # Also check if the queue is being consumed (workers connected)
            info = r.info("clients")
            connected = int(info.get("connected_clients", 0))
            # If there are bindings and multiple connected clients (Django + workers)
            if bindings and connected > 1:
                worker_ok = True
                worker_count = max(1, connected - 1)  # subtract 1 for non-worker connections
        except Exception:
            pass

    core_ok = db_ok and redis_ok
    all_ok = core_ok and worker_ok
    status_code = 200 if core_ok else 503
    return Response(
        {
            "status": "ok" if all_ok else "degraded",
            "timestamp": datetime.datetime.utcnow().isoformat() + "Z",
            "db": db_ok,
            "redis": redis_ok,
            "worker": worker_ok,
            "worker_count": worker_count,
        },
        status=status_code,
    )
